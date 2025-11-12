import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import win32com.client as win32  
import os
# 获取当前目录
current_dir = os.getcwd()


# 定义默认值
default_inspector = "蒋玉长"
default_improvement_suggestions = "在客户端与服务端通信时实现双向认证"

# 读取四次输入内容，带有默认值提示
print(f"检测人员（默认: {default_inspector}):", end="")
inspector = input()
if not inspector:
    inspector = default_inspector

print(f"被测客户端软件概述:", end="")
software_summary = input()

print(f"完善建议（默认: {default_improvement_suggestions}):", end="")
improvement_suggestions = input()
if not improvement_suggestions:
    improvement_suggestions = default_improvement_suggestions

print(f"请输入MD5值:", end="")
MD5 = input()
MD5 = MD5.upper()
# 构建文件路径
file_path = os.path.join(current_dir, '05客户端软件检测总体情况报告.docx')

# 检查文件是否存在
if not os.path.exists(file_path):
    print(f"文件 {file_path} 不存在，请检查文件路径。")
else:
    try:
        # 打开文档
        doc = Document(file_path)

        # 用于记录替换次数
        replace_count = 0

        # 遍历文档中的段落
        for paragraph in doc.paragraphs:
            # 替换段落中的文本
            if '请人工填写' in paragraph.text:
                if replace_count == 0:
                    original_text = paragraph.text
                    new_text = original_text.replace('请人工填写', improvement_suggestions, 1)
                    paragraph.text = ""  # 清空段落内容
                    run = paragraph.add_run(new_text)
                    run.font.size = Pt(12)  # 设置字体大小
                    replace_count += 1
            if '请找被测评方索取' in paragraph.text:
                original_text = paragraph.text
                new_text = original_text.replace('请找被测评方索取', software_summary)
                paragraph.text = ""  # 清空段落内容
                run = paragraph.add_run(new_text)
                run.font.size = Pt(12)  # 设置字体大小
            if '登录、修改密码、交易、个人信息保护。' in paragraph.text:
                original_text = paragraph.text
                new_text = original_text.replace('登录、修改密码、交易、个人信息保护。', original_text)
                paragraph.text = ""  # 清空段落内容
                run = paragraph.add_run(new_text)
                run.font.size = Pt(12)  # 设置字体大小
            if '检测地点为北京市石景山区实兴大街30号院18号楼。' in paragraph.text:
                original_text = paragraph.text
                insert_index = len(original_text) - 25         # 插入MD5值
                new_text = original_text[:insert_index] + '该软件MD5为' + MD5 +'。' + "检测地点为北京市石景山区实兴大街30号院18号楼。"
                paragraph.text = ""
                run = paragraph.add_run(new_text)
                run.font.name = '仿宋'  # 设置字体为仿宋
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(16)  # 设置字号为三号（16磅）
            if '建议结合国家监管政策和行业特点，在客户端与服务端通信时实现双向认证等方面不断提高能力水平。' in paragraph.text:
                original_text = paragraph.text
                new_text = original_text.replace('建议结合国家监管政策和行业特点，在客户端与服务端通信时实现双向认证等方面不断提高能力水平。', ('建议结合国家监管政策和行业特点，'+improvement_suggestions +'等方面不断提高能力水平。'))
                paragraph.text = ""  # 清空段落内容
                run = paragraph.add_run(new_text)
                run.font.name = '仿宋'  # 设置字体为仿宋
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(16)  # 设置字号为三号（16磅）
            # 去除段落中文字的突出显示
            for run in paragraph.runs:
                run.highlight_color = None

        # 遍历文档中的表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 替换表格单元格段落中的文本
                        if '请人工填写' in paragraph.text:
                            if replace_count == 1:
                                original_text = paragraph.text
                                new_text = original_text.replace('请人工填写', inspector, 1)
                                paragraph.text = ""  # 清空段落内容
                                run = paragraph.add_run(new_text)
                                run.font.size = Pt(10.5)  # 设置字体大小
                                replace_count += 1

                        # 去除表格单元格段落中文字的突出显示
                        for run in paragraph.runs:
                            run.highlight_color = None

        # 保存修改后的文档
        doc.save(file_path)

    except Exception as e:
        print(f"处理文件时出现错误: {e}")